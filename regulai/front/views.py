from django.shortcuts import render, redirect, get_object_or_404
from django.conf import settings
from django.db.models import Count, Q
from django.db.models.functions import TruncDate
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required, user_passes_test
from django.core.exceptions import PermissionDenied
from functools import wraps
from django.utils import timezone

from front.models import *
from front.connectors import get_connector, ChatMessage, ChatResponse
from front.rate_limit import (
    login_rate_limit, chat_rate_limit, signup_rate_limit,
    api_rate_limit, rate_limit
)
import yara, io
import json
import base64
import hashlib
import re
import requests
import PyPDF2
import docx
import openpyxl
try:
    import pytesseract
    from PIL import Image
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False


# ==================== PERMISSION DECORATORS ====================

def security_required(view_func):
    """
    Decorator that requires the user to be staff or superuser.
    To access the Security sections (Dashboard, Prompts, Rules, Alerts, Audit).
    """
    @wraps(view_func)
    def _wrapped_view(request, *args, **kwargs):
        if not request.user.is_authenticated:
            from django.contrib.auth.views import redirect_to_login
            return redirect_to_login(request.get_full_path())
        if not (request.user.is_staff or request.user.is_superuser):
            raise PermissionDenied("You do not have permission to access this section.")
        return view_func(request, *args, **kwargs)
    return _wrapped_view


def admin_required(view_func):
    """
    Decorator that requires the user to be superuser.
    To access the Settings/Admin sections.
    """
    @wraps(view_func)
    def _wrapped_view(request, *args, **kwargs):
        if not request.user.is_authenticated:
            from django.contrib.auth.views import redirect_to_login
            return redirect_to_login(request.get_full_path())
        if not request.user.is_superuser:
            raise PermissionDenied("Only administrators can access this section.")
        return view_func(request, *args, **kwargs)
    return _wrapped_view


# ==================== HELPER FUNCTIONS ====================

def get_user_company(user):
    """
    Gets the user's company safely.
    Returns the user's company or Default Company if not set.
    """
    if hasattr(user, 'profile') and user.profile and user.profile.company:
        return user.profile.company
    # Return Default Company as fallback
    default_company, _ = Company.objects.get_or_create(name="Default Company")
    # Ensure default company has a default rules group
    RulesGroup.objects.get_or_create(
        name='Default',
        company=default_company,
        defaults={'description': 'Default rules group'}
    )
    return default_company


def log_action(user, action, details=None, request=None):
    """
    Logs an action in the AuditLog.

    :param user: User performing the action
    :param action: Action type (e.g.: 'login', 'rule_create', 'alert_resolve')
    :param details: Additional details (string or dict)
    :param request: Optional request to obtain additional info (IP, etc)
    """
    if isinstance(details, dict):
        # Add request information if available
        if request:
            details['ip_address'] = get_client_ip(request)
            details['user_agent'] = request.META.get('HTTP_USER_AGENT', '')[:200]
        details = json.dumps(details, ensure_ascii=False, default=str)

    AuditLog.objects.create(
        user=user if user and user.is_authenticated else None,
        action=action,
        details=details
    )


def get_client_ip(request):
    """Gets the client IP address."""
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    if x_forwarded_for:
        ip = x_forwarded_for.split(',')[0]
    else:
        ip = request.META.get('REMOTE_ADDR')
    return ip




# ==================== HOME VIEW ====================

@login_required
def home_view(request):
    """
    Home view with role-based content.
    - Regular users: See chat access and basic info
    - Staff users: See security features
    - Admin users: See full admin capabilities
    """
    user = request.user
    context = {
        'is_admin': user.is_superuser,
        'is_staff': user.is_staff,
        'is_regular': not user.is_staff and not user.is_superuser,
    }

    # Stats for staff/admin users
    if user.is_staff or user.is_superuser:
        user_company = get_user_company(user)
        context['stats'] = {
            'total_prompts': Prompt.objects.filter(user__profile__company=user_company).count(),
            'blocked_prompts': Prompt.objects.filter(user__profile__company=user_company, filtered=True).count(),
            'active_rules': Rule.objects.filter(
                Q(rules_group__company=user_company) | Q(rules_group__company__isnull=True),
                active=True
            ).count(),
            'pending_alerts': Alert.objects.filter(company=user_company, resolved=False).count(),
        }

    # Additional stats for admin users
    if user.is_superuser:
        user_company = get_user_company(user)
        context['admin_stats'] = {
            'total_users': Profile.objects.filter(company=user_company).count(),
            'total_departments': Department.objects.filter(company=user_company).count(),
            'active_engines': CompanyEngine.objects.filter(company=user_company, active=True).count(),
        }

    return render(request, 'home.html', context)


# ==================== AUTH VIEWS ====================

@login_rate_limit
def login_view(request):
    """Login view with rate limiting (5 attempts per minute per IP)."""
    if request.user.is_authenticated:
        return redirect('front:home')

    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        remember = request.POST.get('remember')

        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)

            # Registrar login exitoso
            log_action(user, 'user_login', {
                'username': user.username,
                'remember_me': bool(remember)
            }, request)

            # If "remember me" was not checked, session expires when browser closes
            if not remember:
                request.session.set_expiry(0)

            # Redirect to requested page or home
            next_url = request.POST.get('next') or request.GET.get('next')
            if not next_url:
                next_url = 'front:home'
            return redirect(next_url)
        else:
            # Log failed attempt
            log_action(None, 'user_login_failed', {
                'username': username,
                'reason': 'Invalid credentials'
            }, request)
            messages.error(request, 'Invalid username or password.')

    return render(request, 'auth/login.html', {
        'next': request.GET.get('next', ''),
        'allow_registration': getattr(settings, 'ALLOW_REGISTRATION', True)
    })


def logout_view(request):
    """Logout view."""
    # Log logout before closing session
    if request.user.is_authenticated:
        log_action(request.user, 'user_logout', {
            'username': request.user.username
        }, request)
    logout(request)
    messages.success(request, 'You have been logged out successfully.')
    return redirect('front:login')


@signup_rate_limit
def signup_view(request):
    """
    Public signup view for creating new companies and admin users.
    Rate limited to 3 registrations per hour per IP.

    SECURITY: Only the first user becomes a superuser (platform owner).
    Subsequent users are company admins with staff privileges but NOT superusers.
    """
    # Check if registration is allowed
    if not getattr(settings, 'ALLOW_REGISTRATION', True):
        messages.error(request, 'Registration is currently disabled.')
        return redirect('front:login')

    # Redirect if already authenticated
    if request.user.is_authenticated:
        return redirect('front:home')

    if request.method == 'POST':
        # Get form data
        username = request.POST.get('username', '').strip()
        email = request.POST.get('email', '').strip()
        password = request.POST.get('password', '').strip()
        password_confirm = request.POST.get('password_confirm', '').strip()
        first_name = request.POST.get('first_name', '').strip()
        last_name = request.POST.get('last_name', '').strip()
        company_name = request.POST.get('company_name', '').strip()
        company_domain = request.POST.get('company_domain', '').strip()

        # Validation
        errors = []

        if not username:
            errors.append('Username is required.')
        elif User.objects.filter(username=username).exists():
            errors.append('Username already exists.')

        if not email:
            errors.append('Email is required.')
        elif User.objects.filter(email=email).exists():
            errors.append('Email already exists.')

        if not password:
            errors.append('Password is required.')
        elif len(password) < 8:
            errors.append('Password must be at least 8 characters long.')

        if password != password_confirm:
            errors.append('Passwords do not match.')

        if not company_name:
            errors.append('Company name is required.')
        elif Company.objects.filter(name=company_name).exists():
            errors.append('Company name already exists.')

        if errors:
            for error in errors:
                messages.error(request, error)
            return render(request, 'auth/signup.html', {
                'form_data': request.POST
            })

        try:
            # Create company
            company = Company.objects.create(
                name=company_name,
                domain=company_domain if company_domain else None
            )

            # SECURITY: Only the first user (platform owner) gets superuser privileges
            # Subsequent users are company admins with staff access but NOT superusers
            is_first_user = User.objects.count() == 0

            # Create user with appropriate privileges
            user = User.objects.create_user(
                username=username,
                email=email,
                password=password,
                first_name=first_name,
                last_name=last_name,
                is_staff=True,  # Allow access to security features
                is_superuser=is_first_user  # Only first user is platform superuser
            )

            # Get or create admin role
            admin_role, _ = Role.objects.get_or_create(
                name='Admin',
                defaults={'description': 'Company administrator with full access'}
            )

            # Create default rules group for the company
            RulesGroup.objects.get_or_create(
                name='Default',
                company=company,
                defaults={'description': 'Default rules group'}
            )

            # Create profile
            Profile.objects.create(
                user=user,
                company=company,
                role=admin_role
            )

            # Log the registration
            log_action(user, 'user_registered', {
                'username': username,
                'company': company_name,
                'is_first_user': is_first_user,
                'is_superuser': is_first_user
            }, request)

            # Auto-login the user
            login(request, user)

            # Show appropriate message
            if is_first_user:
                messages.success(request,
                    f'Welcome to regulAI! You are the platform owner. '
                    f'Your company "{company_name}" has been created with full system access.'
                )
            else:
                messages.success(request,
                    f'Welcome to regulAI! Your company "{company_name}" has been created. '
                    f'You have administrator access for your company.'
                )

            return redirect('front:dashboard')

        except Exception as e:
            messages.error(request, f'An error occurred during registration: {str(e)}')
            # Rollback: delete company if user creation failed
            if 'company' in locals():
                company.delete()
            return render(request, 'auth/signup.html', {
                'form_data': request.POST
            })

    return render(request, 'auth/signup.html')


@login_required
def profile_view(request):
    """User profile view."""
    user = request.user

    # Get or create profile
    profile = None
    departments = []
    available_engines = []
    if hasattr(user, 'profile'):
        profile = user.profile
        # Get departments for user's company
        if profile.company:
            departments = Department.objects.filter(company=profile.company).order_by('name')
            # Get available engines for user's company
            available_engines = CompanyEngine.objects.filter(
                company=profile.company,
                active=True
            ).select_related('engine').order_by('engine__name')

    if request.method == 'POST':
        # Update user data
        user.first_name = request.POST.get('first_name', '')
        user.last_name = request.POST.get('last_name', '')
        user.email = request.POST.get('email', '')
        user.save()

        # Update profile if exists
        if profile:
            # Handle department selection
            department_id = request.POST.get('department')
            if department_id:
                # Validate department belongs to user's company
                department = Department.objects.filter(
                    id=department_id,
                    company=profile.company
                ).first()
                profile.department = department
            else:
                profile.department = None

            # Handle avatar upload
            if 'avatar' in request.FILES:
                avatar_file = request.FILES['avatar']
                # Validate file type
                allowed_types = ['image/jpeg', 'image/png', 'image/gif', 'image/webp']
                if avatar_file.content_type in allowed_types:
                    # Delete old avatar if exists
                    if profile.avatar:
                        profile.avatar.delete(save=False)
                    profile.avatar = avatar_file
                else:
                    messages.error(request, 'Invalid image format. Please use JPEG, PNG, GIF or WebP.')
                    return redirect('front:profile')

            # Handle avatar removal
            if request.POST.get('remove_avatar') == 'true' and profile.avatar:
                profile.avatar.delete(save=False)
                profile.avatar = None

            # Handle default engine selection
            default_engine_id = request.POST.get('default_engine')
            if default_engine_id:
                # Validate engine belongs to user's company and is active
                engine = CompanyEngine.objects.filter(
                    id=default_engine_id,
                    company=profile.company,
                    active=True
                ).first()
                profile.default_engine = engine
            else:
                profile.default_engine = None

            profile.save()

        # Change password if provided
        new_password = request.POST.get('new_password')
        confirm_password = request.POST.get('confirm_password')

        if new_password:
            if new_password == confirm_password:
                user.set_password(new_password)
                user.save()
                # Re-authenticate the user
                login(request, user)
                log_action(user, 'profile_update', {
                    'changes': ['profile_data', 'password']
                }, request)
                messages.success(request, 'Profile and password updated successfully.')
            else:
                messages.error(request, 'Passwords do not match.')
                return redirect('front:profile')
        else:
            log_action(user, 'profile_update', {
                'changes': ['profile_data']
            }, request)
            messages.success(request, 'Profile updated successfully.')

        return redirect('front:profile')

    context = {
        'profile': profile,
        'departments': departments,
        'available_engines': available_engines,
    }
    return render(request, 'auth/profile.html', context)


# ==================== USER MANAGEMENT VIEWS ====================

@admin_required
def users_list(request):
    """List users from the admin's company."""
    company = request.user.profile.company
    profiles = Profile.objects.filter(company=company).select_related('user', 'role', 'department').order_by('user__username')

    # Search filter
    search = request.GET.get('search', '')
    if search:
        profiles = profiles.filter(
            Q(user__username__icontains=search) |
            Q(user__email__icontains=search) |
            Q(user__first_name__icontains=search) |
            Q(user__last_name__icontains=search)
        )

    # Role filter
    role_id = request.GET.get('role')
    if role_id:
        profiles = profiles.filter(role_id=role_id)

    # Department filter
    department_id = request.GET.get('department')
    if department_id:
        profiles = profiles.filter(department_id=department_id)

    # Get available roles and departments for filters
    roles = Role.objects.all()
    departments = Department.objects.filter(company=company)

    context = {
        'profiles': profiles,
        'roles': roles,
        'departments': departments,
        'search': search,
        'current_role': role_id,
        'current_department': department_id,
        'total_users': profiles.count(),
    }
    return render(request, 'users/list.html', context)


@admin_required
def user_create(request):
    """Create a new user in the admin's company."""
    company = request.user.profile.company
    roles = Role.objects.all()
    departments = Department.objects.filter(company=company)

    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        email = request.POST.get('email', '').strip()
        password = request.POST.get('password', '')
        confirm_password = request.POST.get('confirm_password', '')
        first_name = request.POST.get('first_name', '').strip()
        last_name = request.POST.get('last_name', '').strip()
        role_id = request.POST.get('role')
        department_id = request.POST.get('department')
        is_staff = request.POST.get('is_staff') == 'on'

        # Validation
        errors = []
        if not username:
            errors.append('Username is required.')
        elif User.objects.filter(username=username).exists():
            errors.append('Username already exists.')

        if not email:
            errors.append('Email is required.')
        elif User.objects.filter(email=email).exists():
            errors.append('Email already exists.')

        if not password:
            errors.append('Password is required.')
        elif password != confirm_password:
            errors.append('Passwords do not match.')

        if errors:
            for error in errors:
                messages.error(request, error)
            return render(request, 'users/form.html', {
                'roles': roles,
                'departments': departments,
                'form_data': request.POST,
                'is_edit': False,
            })

        try:
            # Create user
            user = User.objects.create_user(
                username=username,
                email=email,
                password=password,
                first_name=first_name,
                last_name=last_name,
                is_staff=is_staff,
            )

            # Create profile
            role = Role.objects.filter(id=role_id).first() if role_id else None
            department = Department.objects.filter(id=department_id, company=company).first() if department_id else None
            Profile.objects.create(
                user=user,
                company=company,
                role=role,
                department=department,
            )

            log_action(request.user, 'user_create', {
                'created_user': username,
                'company': company.name,
                'is_staff': is_staff,
            }, request)

            messages.success(request, f'User "{username}" created successfully.')
            return redirect('front:users_list')

        except Exception as e:
            messages.error(request, f'Error creating user: {str(e)}')

    context = {
        'roles': roles,
        'departments': departments,
        'is_edit': False,
    }
    return render(request, 'users/form.html', context)


@admin_required
def user_edit(request, pk):
    """Edit an existing user (only from admin's company)."""
    company = request.user.profile.company
    profile = get_object_or_404(Profile, pk=pk, company=company)
    user = profile.user
    roles = Role.objects.all()
    departments = Department.objects.filter(company=company)

    if request.method == 'POST':
        email = request.POST.get('email', '').strip()
        first_name = request.POST.get('first_name', '').strip()
        last_name = request.POST.get('last_name', '').strip()
        role_id = request.POST.get('role')
        department_id = request.POST.get('department')
        is_staff = request.POST.get('is_staff') == 'on'
        new_password = request.POST.get('password', '')
        confirm_password = request.POST.get('confirm_password', '')

        # Validation
        errors = []
        if not email:
            errors.append('Email is required.')
        elif User.objects.filter(email=email).exclude(pk=user.pk).exists():
            errors.append('Email already exists.')

        if new_password and new_password != confirm_password:
            errors.append('Passwords do not match.')

        if errors:
            for error in errors:
                messages.error(request, error)
            return render(request, 'users/form.html', {
                'roles': roles,
                'departments': departments,
                'profile': profile,
                'edit_user': user,
                'is_edit': True,
            })

        # Update user
        user.email = email
        user.first_name = first_name
        user.last_name = last_name
        user.is_staff = is_staff
        if new_password:
            user.set_password(new_password)
        user.save()

        # Update profile
        profile.role = Role.objects.filter(id=role_id).first() if role_id else None
        profile.department = Department.objects.filter(id=department_id, company=company).first() if department_id else None
        profile.save()

        log_action(request.user, 'user_edit', {
            'edited_user': user.username,
            'changes': ['profile_data'] + (['password'] if new_password else []),
        }, request)

        messages.success(request, f'User "{user.username}" updated successfully.')
        return redirect('front:users_list')

    context = {
        'roles': roles,
        'departments': departments,
        'profile': profile,
        'edit_user': user,
        'is_edit': True,
    }
    return render(request, 'users/form.html', context)


@admin_required
def user_delete(request, pk):
    """Delete a user (only from admin's company)."""
    company = request.user.profile.company
    profile = get_object_or_404(Profile, pk=pk, company=company)
    user = profile.user

    # Prevent self-deletion
    if user == request.user:
        messages.error(request, 'You cannot delete your own account.')
        return redirect('front:users_list')

    if request.method == 'POST':
        username = user.username
        user.delete()  # This will cascade delete the profile

        log_action(request.user, 'user_delete', {
            'deleted_user': username,
            'company': company.name,
        }, request)

        messages.success(request, f'User "{username}" deleted successfully.')
        return redirect('front:users_list')

    context = {
        'profile': profile,
        'delete_user': user,
    }
    return render(request, 'users/delete.html', context)


# ==================== DEPARTMENT MANAGEMENT VIEWS ====================

@admin_required
def departments_list(request):
    """List departments from the admin's company."""
    company = request.user.profile.company
    departments = Department.objects.filter(company=company).annotate(
        users_count=Count('users')
    ).order_by('name')

    context = {
        'departments': departments,
        'total_departments': departments.count(),
    }
    return render(request, 'departments/list.html', context)


@admin_required
def department_create(request):
    """Create a new department in the admin's company."""
    company = request.user.profile.company

    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        description = request.POST.get('description', '').strip()

        if not name:
            messages.error(request, 'Department name is required.')
            return render(request, 'departments/form.html', {
                'form_data': request.POST,
                'is_edit': False,
            })

        if Department.objects.filter(name=name, company=company).exists():
            messages.error(request, 'A department with this name already exists.')
            return render(request, 'departments/form.html', {
                'form_data': request.POST,
                'is_edit': False,
            })

        Department.objects.create(
            name=name,
            company=company,
            description=description,
        )

        log_action(request.user, 'department_create', {
            'department_name': name,
            'company': company.name,
        }, request)

        messages.success(request, f'Department "{name}" created successfully.')
        return redirect('front:departments_list')

    return render(request, 'departments/form.html', {'is_edit': False})


@admin_required
def department_edit(request, pk):
    """Edit an existing department (only from admin's company)."""
    company = request.user.profile.company
    department = get_object_or_404(Department, pk=pk, company=company)

    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        description = request.POST.get('description', '').strip()

        if not name:
            messages.error(request, 'Department name is required.')
            return render(request, 'departments/form.html', {
                'department': department,
                'is_edit': True,
            })

        if Department.objects.filter(name=name, company=company).exclude(pk=pk).exists():
            messages.error(request, 'A department with this name already exists.')
            return render(request, 'departments/form.html', {
                'department': department,
                'is_edit': True,
            })

        department.name = name
        department.description = description
        department.save()

        log_action(request.user, 'department_edit', {
            'department_name': name,
        }, request)

        messages.success(request, f'Department "{name}" updated successfully.')
        return redirect('front:departments_list')

    context = {
        'department': department,
        'is_edit': True,
    }
    return render(request, 'departments/form.html', context)


@admin_required
def department_delete(request, pk):
    """Delete a department (only from admin's company)."""
    company = request.user.profile.company
    department = get_object_or_404(Department, pk=pk, company=company)

    if request.method == 'POST':
        name = department.name
        department.delete()

        log_action(request.user, 'department_delete', {
            'department_name': name,
            'company': company.name,
        }, request)

        messages.success(request, f'Department "{name}" deleted successfully.')
        return redirect('front:departments_list')

    context = {
        'department': department,
    }
    return render(request, 'departments/delete.html', context)


# ==================== COMPANY ENGINES VIEWS ====================

@admin_required
def engines_list(request):
    """List all engines configured for the admin's company."""
    company = request.user.profile.company
    company_engines = CompanyEngine.objects.filter(company=company).select_related('engine').prefetch_related('engine__models').order_by('engine__name')

    # Count total models across all company engines
    total_models = sum(ce.engine.models.count() for ce in company_engines)

    context = {
        'company_engines': company_engines,
        'total_engines': company_engines.count(),
        'active_engines': company_engines.filter(active=True).count(),
        'total_models': total_models,
    }
    return render(request, 'engines/list.html', context)


@admin_required
def engine_create(request):
    """Add a new engine to the company."""
    company = request.user.profile.company

    # Get all available engines
    available_engines = Engine.objects.all().order_by('name')

    if request.method == 'POST':
        engine_id = request.POST.get('engine')
        name = request.POST.get('name', '').strip()
        api_key = request.POST.get('api_key', '').strip()
        active = request.POST.get('active') == 'on'

        # Validation
        errors = []
        if not engine_id:
            errors.append('Please select an engine.')
        if not api_key:
            errors.append('API Key is required.')

        engine = Engine.objects.filter(id=engine_id).first()
        if not engine:
            errors.append('Invalid engine selected.')

        if errors:
            for error in errors:
                messages.error(request, error)
            return render(request, 'engines/form.html', {
                'available_engines': available_engines,
                'form_data': request.POST,
                'is_edit': False,
            })

        CompanyEngine.objects.create(
            company=company,
            engine=engine,
            name=name,
            api_key=api_key,
            active=active,
        )

        display_name = name if name else engine.name
        log_action(request.user, 'engine_create', {
            'engine_name': engine.name,
            'display_name': display_name,
            'company': company.name,
            'active': active,
        }, request)

        messages.success(request, f'Engine "{display_name}" added successfully.')
        return redirect('front:engines_list')

    context = {
        'available_engines': available_engines,
        'is_edit': False,
    }
    return render(request, 'engines/form.html', context)


@admin_required
def engine_edit(request, pk):
    """Edit an existing company engine configuration."""
    company = request.user.profile.company
    company_engine = get_object_or_404(CompanyEngine, pk=pk, company=company)

    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        api_key = request.POST.get('api_key', '').strip()
        active = request.POST.get('active') == 'on'

        # Validation
        if not api_key:
            messages.error(request, 'API Key is required.')
            return render(request, 'engines/form.html', {
                'company_engine': company_engine,
                'is_edit': True,
            })

        company_engine.name = name
        company_engine.api_key = api_key
        company_engine.active = active
        company_engine.save()

        log_action(request.user, 'engine_edit', {
            'engine_name': company_engine.engine.name,
            'display_name': company_engine.display_name,
            'company': company.name,
            'active': active,
        }, request)

        messages.success(request, f'Engine "{company_engine.display_name}" updated successfully.')
        return redirect('front:engines_list')

    context = {
        'company_engine': company_engine,
        'is_edit': True,
    }
    return render(request, 'engines/form.html', context)


@admin_required
def engine_delete(request, pk):
    """Remove an engine from the company."""
    company = request.user.profile.company
    company_engine = get_object_or_404(CompanyEngine, pk=pk, company=company)

    if request.method == 'POST':
        engine_name = company_engine.engine.name
        company_engine.delete()

        log_action(request.user, 'engine_delete', {
            'engine_name': engine_name,
            'company': company.name,
        }, request)

        messages.success(request, f'Engine "{engine_name}" removed successfully.')
        return redirect('front:engines_list')

    context = {
        'company_engine': company_engine,
    }
    return render(request, 'engines/delete.html', context)


@admin_required
def engine_toggle(request, pk):
    """Toggle engine active status."""
    company = request.user.profile.company
    company_engine = get_object_or_404(CompanyEngine, pk=pk, company=company)

    company_engine.active = not company_engine.active
    company_engine.save()

    status = 'activated' if company_engine.active else 'deactivated'
    log_action(request.user, 'engine_toggle', {
        'engine_name': company_engine.engine.name,
        'new_status': status,
    }, request)

    messages.success(request, f'Engine "{company_engine.engine.name}" {status}.')
    return redirect('front:engines_list')


# ==================== ENGINE MODELS VIEWS ====================

@admin_required
def model_create(request, engine_pk):
    """Create a new model for an engine."""
    engine = get_object_or_404(Engine, pk=engine_pk)

    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        model_id = request.POST.get('model_id', '').strip()
        description = request.POST.get('description', '').strip()
        max_tokens = request.POST.get('max_tokens', 4096)
        order = request.POST.get('order', 0)
        supports_vision = request.POST.get('supports_vision') == 'on'
        active = request.POST.get('active') == 'on'

        if not name or not model_id:
            messages.error(request, 'Model name and ID are required.')
            return redirect('front:engines_list')

        # Check if model_id already exists for this engine
        if EngineModel.objects.filter(engine=engine, model_id=model_id).exists():
            messages.error(request, f'Model ID "{model_id}" already exists for this engine.')
            return redirect('front:engines_list')

        EngineModel.objects.create(
            engine=engine,
            name=name,
            model_id=model_id,
            description=description,
            max_tokens=int(max_tokens) if max_tokens else 4096,
            order=int(order) if order else 0,
            supports_vision=supports_vision,
            active=active,
        )

        log_action(request.user, 'model_create', {
            'engine_name': engine.name,
            'model_name': name,
            'model_id': model_id,
        }, request)

        messages.success(request, f'Model "{name}" added to {engine.name}.')

    return redirect('front:engines_list')


@admin_required
def model_edit(request, pk):
    """Edit an existing model."""
    model = get_object_or_404(EngineModel, pk=pk)

    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        model_id = request.POST.get('model_id', '').strip()
        description = request.POST.get('description', '').strip()
        max_tokens = request.POST.get('max_tokens', 4096)
        order = request.POST.get('order', 0)
        supports_vision = request.POST.get('supports_vision') == 'on'
        active = request.POST.get('active') == 'on'

        if not name or not model_id:
            messages.error(request, 'Model name and ID are required.')
            return redirect('front:engines_list')

        # Check if model_id already exists for this engine (excluding current model)
        if EngineModel.objects.filter(engine=model.engine, model_id=model_id).exclude(pk=pk).exists():
            messages.error(request, f'Model ID "{model_id}" already exists for this engine.')
            return redirect('front:engines_list')

        model.name = name
        model.model_id = model_id
        model.description = description
        model.max_tokens = int(max_tokens) if max_tokens else 4096
        model.order = int(order) if order else 0
        model.supports_vision = supports_vision
        model.active = active
        model.save()

        log_action(request.user, 'model_edit', {
            'engine_name': model.engine.name,
            'model_name': name,
            'model_id': model_id,
        }, request)

        messages.success(request, f'Model "{name}" updated.')

    return redirect('front:engines_list')


@admin_required
def model_toggle(request, pk):
    """Toggle model active status."""
    model = get_object_or_404(EngineModel, pk=pk)

    model.active = not model.active
    model.save()

    status = 'activated' if model.active else 'deactivated'
    log_action(request.user, 'model_toggle', {
        'engine_name': model.engine.name,
        'model_name': model.name,
        'new_status': status,
    }, request)

    messages.success(request, f'Model "{model.name}" {status}.')
    return redirect('front:engines_list')


@admin_required
def model_delete(request, pk):
    """Delete a model."""
    model = get_object_or_404(EngineModel, pk=pk)
    model_name = model.name
    engine_name = model.engine.name

    model.delete()

    log_action(request.user, 'model_delete', {
        'engine_name': engine_name,
        'model_name': model_name,
    }, request)

    messages.success(request, f'Model "{model_name}" deleted from {engine_name}.')
    return redirect('front:engines_list')


# ==================== DASHBOARD & CHAT VIEWS ====================

@security_required
def dashboard_view(request):
    """Main dashboard with system statistics."""
    from datetime import timedelta

    # Get user's company
    user_company = get_user_company(request.user)

    # Basic statistics - filtered by company
    total_prompts = Prompt.objects.filter(user__profile__company=user_company).count()
    total_alerts = Alert.objects.filter(company=user_company).count()
    # Rules: company rules + global rules (company=null)
    total_rules = Rule.objects.filter(
        active=True
    ).filter(
        Q(rules_group__company=user_company) | Q(rules_group__company__isnull=True)
    ).count()
    total_users = User.objects.filter(profile__company=user_company).count()

    # Recent alerts - filtered by company
    recent_alerts = Alert.objects.filter(
        company=user_company
    ).select_related('user', 'rule').order_by('-created_at')[:5]

    # Recent prompts - filtered by company
    recent_prompts = Prompt.objects.filter(
        user__profile__company=user_company
    ).select_related('user').order_by('-created_at')[:5]

    # Alerts by severity - filtered by company
    alerts_by_severity = Alert.objects.filter(
        company=user_company
    ).values('severity').annotate(count=Count('id'))

    # Alerts timeline for the last 5 days (for bar chart) - filtered by company
    five_days_ago = timezone.now() - timedelta(days=5)
    alerts_timeline = Alert.objects.filter(
        company=user_company,
        created_at__gte=five_days_ago
    ).annotate(
        date=TruncDate('created_at')
    ).values('date', 'severity').annotate(
        count=Count('id')
    ).order_by('date')

    # Convert to list for JSON serialization in template
    alerts_timeline_data = list(alerts_timeline)
    for item in alerts_timeline_data:
        item['date'] = item['date'].isoformat() if item['date'] else None

    # Prompts timeline for the last 5 days (for line overlay on bar chart)
    prompts_timeline = Prompt.objects.filter(
        user__profile__company=user_company,
        created_at__gte=five_days_ago
    ).annotate(
        date=TruncDate('created_at')
    ).values('date').annotate(
        count=Count('id')
    ).order_by('date')

    prompts_timeline_data = list(prompts_timeline)
    for item in prompts_timeline_data:
        item['date'] = item['date'].isoformat() if item['date'] else None

    context = {
        'total_prompts': total_prompts,
        'total_alerts': total_alerts,
        'total_rules': total_rules,
        'total_users': total_users,
        'recent_alerts': recent_alerts,
        'recent_prompts': recent_prompts,
        'alerts_by_severity': alerts_by_severity,
        'alerts_timeline_data': json.dumps(alerts_timeline_data),
        'prompts_timeline_data': json.dumps(prompts_timeline_data),
    }
    return render(request, 'dashboard.html', context)


@login_required
def chat_view(request, conversation_id=None):
    """Chat view with conversation support."""
    user = request.user
    blocked = False
    blocked_rules = []

    # Ensure Profile and Company exist
    if not hasattr(user, 'profile'):
        default_company, _ = Company.objects.get_or_create(name="Default Company")
        Profile.objects.create(user=user, company=default_company)
        user.refresh_from_db()

    profile = user.profile

    # Get available engines for user's company
    available_engines = CompanyEngine.objects.filter(
        company=profile.company,
        active=True
    ).select_related('engine').prefetch_related('engine__models').order_by('engine__name')

    # Determine which engine to use
    # Priority: 1. Session selected, 2. User's default, 3. First available
    selected_engine_id = request.session.get('selected_engine_id')
    current_engine = None

    if selected_engine_id:
        current_engine = available_engines.filter(id=selected_engine_id).first()

    if not current_engine and profile.default_engine:
        current_engine = available_engines.filter(id=profile.default_engine.id).first()

    if not current_engine:
        current_engine = available_engines.first()

    # Determine which model to use
    # Priority: 1. Session selected, 2. User's default (if same engine), 3. Engine's default model
    selected_model_id = request.session.get('selected_model_id')
    current_model = None
    available_models = []

    if current_engine:
        available_models = current_engine.engine.models.filter(active=True).order_by('order', 'name')

        if selected_model_id:
            current_model = available_models.filter(id=selected_model_id).first()

        if not current_model and profile.default_model:
            # Only use default if it belongs to the current engine
            if profile.default_model.engine_id == current_engine.engine_id:
                current_model = available_models.filter(id=profile.default_model.id).first()

        if not current_model:
            current_model = available_models.first()

    # Get vision support from selected model
    supports_vision = False
    if current_model:
        supports_vision = current_model.supports_vision

    # Get API Key and settings from selected engine/model for connector
    api_key = settings.OPENAI_API_KEY
    base_url = None
    model_name = "gpt-4o-mini"
    connector_type = "openai"

    if current_engine:
        api_key = current_engine.api_key
        if current_engine.engine.base_url:
            base_url = current_engine.engine.base_url
        connector_type = current_engine.engine.connector_type or "openai"

        if current_model:
            model_name = current_model.model_id
        elif current_engine.engine.default_model:
            model_name = current_engine.engine.default_model

    # Create connector based on engine type (needed for POST processing)
    connector = None
    if api_key and api_key.strip():
        connector = get_connector(connector_type, api_key, base_url)

    # Get or create conversation
    conversation = None
    if conversation_id:
        conversation = get_object_or_404(Conversation, id=conversation_id, user=user)

    # Get user's conversations for sidebar
    conversations = Conversation.objects.filter(user=user).order_by('-updated_at')[:20]

    if request.method == "POST":
        prompt_text = request.POST.get("prompt", "").strip()
        uploaded_file = request.FILES.get('file')

        # Read file content if uploaded
        file_content_text = None
        file_name = None
        file_size = 0
        file_hashes = None
        is_image = False
        image_base64 = None
        image_mime_type = None

        if uploaded_file:
            file_name = uploaded_file.name
            file_size = uploaded_file.size
            file_extension = file_name.lower().split('.')[-1] if '.' in file_name else ''

            # Check if it's an image
            image_extensions = ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp']
            is_image = file_extension in image_extensions
            mime_types = {
                'png': 'image/png',
                'jpg': 'image/jpeg',
                'jpeg': 'image/jpeg',
                'gif': 'image/gif',
                'bmp': 'image/bmp',
                'webp': 'image/webp'
            }
            if is_image:
                image_mime_type = mime_types.get(file_extension, 'image/png')

            try:
                # Read file content
                file_content = uploaded_file.read()

                # Calculate file hashes
                file_hashes = {
                    'md5': hashlib.md5(file_content).hexdigest(),
                    'sha1': hashlib.sha1(file_content).hexdigest(),
                    'sha256': hashlib.sha256(file_content).hexdigest(),
                }

                # Handle PDF files - extract text
                if file_extension == 'pdf':
                    try:
                        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                        extracted_text = []
                        for page_num, page in enumerate(pdf_reader.pages, 1):
                            page_text = page.extract_text()
                            if page_text:
                                extracted_text.append(f"--- Page {page_num} ---\n{page_text}")
                        file_content_text = "\n\n".join(extracted_text) if extracted_text else "[PDF without extractable text]"
                    except Exception as pdf_error:
                        file_content_text = f"[Error extracting PDF text: {str(pdf_error)}]"

                # Handle DOCX files - extract text
                elif file_extension == 'docx':
                    try:
                        doc = docx.Document(io.BytesIO(file_content))
                        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                        # Also extract text from tables
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                                if row_text:
                                    paragraphs.append(" | ".join(row_text))
                        file_content_text = "\n\n".join(paragraphs) if paragraphs else "[DOCX without extractable text]"
                    except Exception as docx_error:
                        file_content_text = f"[Error extracting DOCX text: {str(docx_error)}]"

                # Handle Excel files (xlsx, xls)
                elif file_extension in ['xlsx', 'xls']:
                    try:
                        workbook = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
                        extracted_text = []
                        for sheet_name in workbook.sheetnames:
                            sheet = workbook[sheet_name]
                            sheet_text = [f"--- Sheet: {sheet_name} ---"]
                            for row in sheet.iter_rows(values_only=True):
                                # Filter out empty cells and convert to strings
                                row_values = [str(cell) if cell is not None else "" for cell in row]
                                if any(v.strip() for v in row_values):
                                    sheet_text.append(" | ".join(row_values))
                            if len(sheet_text) > 1:  # Has content beyond header
                                extracted_text.append("\n".join(sheet_text))
                        file_content_text = "\n\n".join(extracted_text) if extracted_text else "[Excel without extractable data]"
                    except Exception as excel_error:
                        file_content_text = f"[Error extracting Excel data: {str(excel_error)}]"

                # Handle CSV files
                elif file_extension == 'csv':
                    try:
                        file_content_text = file_content.decode('utf-8')
                    except UnicodeDecodeError:
                        file_content_text = file_content.decode('latin-1')

                # Handle image files - OCR extraction + send to vision API
                elif file_extension in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp']:
                    # Store base64 for vision API
                    image_base64 = base64.b64encode(file_content).decode('utf-8')

                    # Extract text via OCR if pytesseract is available (for YARA validation)
                    if PYTESSERACT_AVAILABLE:
                        try:
                            image = Image.open(io.BytesIO(file_content))
                            ocr_text = pytesseract.image_to_string(image)
                            file_content_text = ocr_text.strip() if ocr_text.strip() else "[Image without detected text]"
                        except Exception as ocr_error:
                            file_content_text = f"[OCR error: {str(ocr_error)}]"
                    else:
                        file_content_text = "[Image file]"

                # Handle text-based files
                else:
                    try:
                        file_content_text = file_content.decode('utf-8')
                    except UnicodeDecodeError:
                        try:
                            file_content_text = file_content.decode('latin-1')
                        except:
                            file_content_text = f"[Binary file: {file_name}]"

            except Exception as e:
                file_content_text = f"[Error reading file: {str(e)}]"

        # Need either prompt or file
        if not prompt_text and not uploaded_file:
            return redirect('front:chat_conversation', conversation_id=conversation.id) if conversation else redirect('front:chat')

        # Build the full content for the message
        if file_content_text and prompt_text:
            full_content = f"{prompt_text}\n\n--- Attached File: {file_name} ---\n{file_content_text[:10000]}"
            display_content = prompt_text
        elif file_content_text:
            full_content = f"Please analyze this file:\n\n--- File: {file_name} ---\n{file_content_text[:10000]}"
            display_content = f"[Attached file: {file_name}]"
            prompt_text = display_content
        else:
            full_content = prompt_text
            display_content = prompt_text

        # Create new conversation if needed
        if not conversation:
            title = prompt_text[:50] if prompt_text else file_name[:50] if file_name else "New conversation"
            conversation = Conversation.objects.create(
                user=user,
                title=title + ('...' if len(title) > 50 else '')
            )

        # Create user message
        user_message = Message.objects.create(
            conversation=conversation,
            role='user',
            content=display_content,
            file_name=file_name,
            file_size=file_size
        )

        # Create prompt for audit/tracking
        prompt_obj = Prompt.objects.create(
            user=user,
            conversation=conversation,
            content=prompt_text or f"[File: {file_name}]",
            file_name=file_name,
            file_size=file_size,
            file_md5=file_hashes.get('md5') if file_hashes else None,
            file_sha1=file_hashes.get('sha1') if file_hashes else None,
            file_sha256=file_hashes.get('sha256') if file_hashes else None
        )

        # Validate prompt against rules (for prompts type)
        prompt_matches, prompt_matched_rules = validate_prompt(user, prompt_text, prompt_obj, None, 'prompts')

        # Validate file content against rules (for files type) if file uploaded
        file_matches = []
        file_matched_rules = []
        if file_content_text:
            import logging
            logging.info(f"Validating file content. Length: {len(file_content_text)}, Content preview: {file_content_text[:200]}")
            file_info = {'name': file_name, 'size': file_size, 'hashes': file_hashes} if file_hashes else None
            file_matches, file_matched_rules = validate_prompt(user, file_content_text, prompt_obj, None, 'files', file_info)
            logging.info(f"File validation results: {len(file_matches)} matches, {len(file_matched_rules)} matched rules")

        # Combine matches
        all_matches = (prompt_matches or []) + (file_matches or [])
        all_matched_rules = prompt_matched_rules + file_matched_rules

        if all_matched_rules:
            # Create alerts for all matched rules
            for rule_data in all_matched_rules:
                try:
                    from front.models import Rule
                    rule = Rule.objects.get(id=rule_data['id'])

                    alert_description = (
                        f"Prompt blocked by rule '{rule.name}'. "
                        f"The message was not sent to the LLM."
                    )

                    Alert.objects.create(
                        user=user,
                        company=profile.company,
                        prompt=prompt_obj,
                        rule=rule,
                        severity=rule_data.get('severity', 'medium'),
                        description=alert_description
                    )

                    log_action(user, 'blocking_alert_created', {
                        'rule_name': rule.name,
                        'severity': rule_data.get('severity', 'medium'),
                        'prompt_id': prompt_obj.id,
                        'conversation_id': conversation.id
                    }, request)

                except Rule.DoesNotExist:
                    import logging
                    logging.warning(f"Rule {rule_data.get('name', 'unknown')} not found for alert creation")

            # Mark as blocked
            prompt_obj.filtered = True
            prompt_obj.save()
            user_message.blocked = True
            user_message.save()
            blocked = True
            blocked_rules = all_matched_rules

        # Log the action
        log_action(user, 'prompt_sent', {
            'prompt_id': prompt_obj.id,
            'conversation_id': conversation.id,
            'content_length': len(prompt_text) if prompt_text else 0,
            'has_file': bool(uploaded_file),
            'file_name': file_name,
            'file_size': file_size,
            'rules_matched': len(all_matches) if all_matches else 0,
            'blocked': blocked
        }, request)

        # Only send to LLM if not blocked and connector is available
        if not blocked and connector:
            try:
                # Build messages history for context using ChatMessage objects
                chat_messages = []
                for msg in conversation.messages.exclude(id=user_message.id):
                    if msg.role in ['user', 'assistant']:
                        chat_messages.append(ChatMessage(
                            role=msg.role,
                            content=msg.content
                        ))

                # Add current message - handle images with vision API
                if is_image and image_base64:
                    # Build multimodal content for vision API
                    user_content = []
                    if prompt_text and prompt_text != f"[Attached file: {file_name}]":
                        user_content.append({
                            "type": "text",
                            "text": prompt_text
                        })
                    else:
                        user_content.append({
                            "type": "text",
                            "text": "Please analyze this image."
                        })
                    user_content.append({
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{image_mime_type};base64,{image_base64}"
                        }
                    })
                    chat_messages.append(ChatMessage(
                        role="user",
                        content=user_content
                    ))
                else:
                    # Regular text message
                    chat_messages.append(ChatMessage(
                        role="user",
                        content=full_content
                    ))

                # Use connector to send message
                ai_response = connector.chat(
                    messages=chat_messages,
                    model=model_name,
                    max_tokens=1000,
                    temperature=0.6,
                )

                # Save assistant response as message
                Message.objects.create(
                    conversation=conversation,
                    role='assistant',
                    content=ai_response.content,
                    tokens_used=ai_response.tokens_used
                )

                # Update conversation timestamp
                conversation.save()

            except Exception as e:
                # Save error as system message
                Message.objects.create(
                    conversation=conversation,
                    role='system',
                    content=f"Error: {str(e)}"
                )

        # Redirect to conversation view to avoid form resubmission
        return redirect('front:chat_conversation', conversation_id=conversation.id)

    # Get messages for current conversation
    chat_messages = []
    if conversation:
        chat_messages = list(conversation.messages.all())

    return render(request, "chat/chat.html", {
        "conversation": conversation,
        "conversations": conversations,
        "messages": chat_messages,
        "blocked": blocked,
        "blocked_rules": blocked_rules,
        "available_engines": available_engines,
        "current_engine": current_engine,
        "available_models": available_models,
        "current_model": current_model,
    })


@login_required
def chat_new(request):
    """Creates a new conversation."""
    return redirect('front:chat')


@login_required
def chat_delete(request, conversation_id):
    """Deletes a conversation."""
    conversation = get_object_or_404(Conversation, id=conversation_id, user=request.user)

    if request.method == 'POST':
        log_action(request.user, 'conversation_delete', {
            'conversation_id': conversation.id,
            'title': conversation.title
        }, request)
        conversation.delete()
        messages.success(request, 'Conversation deleted successfully.')
        return redirect('front:chat')

    return render(request, 'chat/delete.html', {'conversation': conversation})


from django.http import JsonResponse
from django.views.decorators.http import require_POST


@login_required
@require_POST
def chat_select_engine(request):
    """API endpoint to select an engine for the chat session."""
    engine_id = request.POST.get('engine_id')

    if engine_id:
        # Validate the engine belongs to user's company
        company = request.user.profile.company
        engine = CompanyEngine.objects.filter(
            id=engine_id,
            company=company,
            active=True
        ).select_related('engine').prefetch_related('engine__models').first()

        if engine:
            request.session['selected_engine_id'] = engine.id
            # Clear model selection when engine changes
            request.session.pop('selected_model_id', None)

            # Get available models for the new engine
            models = list(engine.engine.models.filter(active=True).order_by('order', 'name').values(
                'id', 'name', 'model_id', 'supports_vision'
            ))

            return JsonResponse({
                'success': True,
                'engine_name': engine.engine.name,
                'engine_provider': engine.engine.provider or '',
                'available_models': models,
            })
        else:
            return JsonResponse({
                'success': False,
                'error': 'Invalid engine selected'
            }, status=400)
    else:
        # Clear selection (use default)
        request.session.pop('selected_engine_id', None)
        request.session.pop('selected_model_id', None)
        return JsonResponse({'success': True, 'engine_name': None})


@login_required
@require_POST
def chat_select_model(request):
    """API endpoint to select a model for the chat session."""
    from .models import EngineModel

    model_id = request.POST.get('model_id')

    if model_id:
        # Get current engine from session
        selected_engine_id = request.session.get('selected_engine_id')
        profile = request.user.profile

        # Determine current engine
        current_engine = None
        if selected_engine_id:
            current_engine = CompanyEngine.objects.filter(
                id=selected_engine_id,
                company=profile.company,
                active=True
            ).select_related('engine').first()

        if not current_engine and profile.default_engine:
            current_engine = CompanyEngine.objects.filter(
                id=profile.default_engine.id,
                company=profile.company,
                active=True
            ).select_related('engine').first()

        if not current_engine:
            current_engine = CompanyEngine.objects.filter(
                company=profile.company,
                active=True
            ).select_related('engine').first()

        if not current_engine:
            return JsonResponse({
                'success': False,
                'error': 'No engine available'
            }, status=400)

        # Validate the model belongs to the current engine
        model = EngineModel.objects.filter(
            id=model_id,
            engine=current_engine.engine,
            active=True
        ).first()

        if model:
            request.session['selected_model_id'] = model.id
            return JsonResponse({
                'success': True,
                'model_name': model.name,
                'model_id': model.model_id,
                'supports_vision': model.supports_vision,
            })
        else:
            return JsonResponse({
                'success': False,
                'error': 'Invalid model selected'
            }, status=400)
    else:
        # Clear selection (use default)
        request.session.pop('selected_model_id', None)
        return JsonResponse({'success': True, 'model_name': None})


@login_required
@require_POST
@chat_rate_limit
def chat_send_message(request, conversation_id=None):
    """API endpoint for sending chat messages asynchronously (rate limited to 30 per minute)."""
    user = request.user
    blocked = False
    blocked_rules = []

    # Ensure Profile and Company exist
    if not hasattr(user, 'profile'):
        default_company, _ = Company.objects.get_or_create(name="Default Company")
        Profile.objects.create(user=user, company=default_company)
        user.refresh_from_db()

    profile = user.profile
    company = profile.company

    # Get available engines for user's company
    available_engines = CompanyEngine.objects.filter(
        company=profile.company,
        active=True
    ).select_related('engine').prefetch_related('engine__models')

    # Determine which engine to use (same logic as chat_view)
    # Priority: 1. Session selected, 2. User's default, 3. First available
    selected_engine_id = request.session.get('selected_engine_id')
    current_engine = None

    if selected_engine_id:
        current_engine = available_engines.filter(id=selected_engine_id).first()

    if not current_engine and profile.default_engine:
        current_engine = available_engines.filter(id=profile.default_engine.id).first()

    if not current_engine:
        current_engine = available_engines.first()

    # Determine which model to use
    selected_model_id = request.session.get('selected_model_id')
    current_model = None

    if current_engine:
        available_models = current_engine.engine.models.filter(active=True)

        if selected_model_id:
            current_model = available_models.filter(id=selected_model_id).first()

        if not current_model and profile.default_model:
            if profile.default_model.engine_id == current_engine.engine_id:
                current_model = available_models.filter(id=profile.default_model.id).first()

        if not current_model:
            current_model = available_models.first()

    # Get API Key and settings from selected engine/model
    api_key = settings.OPENAI_API_KEY
    base_url = None
    model_name = "gpt-4o-mini"
    connector_type = "openai"

    if current_engine:
        api_key = current_engine.api_key
        if current_engine.engine.base_url:
            base_url = current_engine.engine.base_url
        connector_type = current_engine.engine.connector_type or "openai"

        if current_model:
            model_name = current_model.model_id
        elif current_engine.engine.default_model:
            model_name = current_engine.engine.default_model

    # Validate API key before proceeding
    if not api_key or api_key.strip() == '':
        return JsonResponse({
            'success': False,
            'error_type': 'no_api_key',
            'error': 'No API key configured for the selected AI engine.',
            'engine_name': current_engine.name if current_engine else 'Default',
        })

    # Create connector based on engine type
    connector = get_connector(connector_type, api_key, base_url)

    # Get or create conversation
    conversation = None
    if conversation_id:
        conversation = get_object_or_404(Conversation, id=conversation_id, user=user)

    prompt_text = request.POST.get("prompt", "").strip()
    uploaded_file = request.FILES.get('file')

    # Read file content if uploaded
    file_content_text = None
    file_name = None
    file_size = 0
    file_hashes = None
    is_image = False
    image_base64 = None
    image_mime_type = None

    if uploaded_file:
        file_name = uploaded_file.name
        file_size = uploaded_file.size
        file_extension = file_name.lower().split('.')[-1] if '.' in file_name else ''

        # Check if it's an image
        image_extensions = ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp']
        is_image = file_extension in image_extensions
        mime_types = {
            'png': 'image/png',
            'jpg': 'image/jpeg',
            'jpeg': 'image/jpeg',
            'gif': 'image/gif',
            'bmp': 'image/bmp',
            'webp': 'image/webp'
        }
        if is_image:
            image_mime_type = mime_types.get(file_extension, 'image/png')

        try:
            file_content = uploaded_file.read()

            # Calculate file hashes
            file_hashes = {
                'md5': hashlib.md5(file_content).hexdigest(),
                'sha1': hashlib.sha1(file_content).hexdigest(),
                'sha256': hashlib.sha256(file_content).hexdigest(),
            }

            if file_extension == 'pdf':
                try:
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                    extracted_text = []
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            extracted_text.append(f"--- Page {page_num} ---\n{page_text}")
                    file_content_text = "\n\n".join(extracted_text) if extracted_text else "[PDF without extractable text]"
                except Exception as pdf_error:
                    file_content_text = f"[Error extracting PDF text: {str(pdf_error)}]"

            elif file_extension == 'docx':
                try:
                    doc = docx.Document(io.BytesIO(file_content))
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_text:
                                paragraphs.append(" | ".join(row_text))
                    file_content_text = "\n\n".join(paragraphs) if paragraphs else "[DOCX without extractable text]"
                except Exception as docx_error:
                    file_content_text = f"[Error extracting DOCX text: {str(docx_error)}]"

            elif file_extension in ['xlsx', 'xls']:
                try:
                    workbook = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
                    extracted_text = []
                    for sheet_name in workbook.sheetnames:
                        sheet = workbook[sheet_name]
                        sheet_text = [f"--- Sheet: {sheet_name} ---"]
                        for row in sheet.iter_rows(values_only=True):
                            row_values = [str(cell) if cell is not None else "" for cell in row]
                            if any(v.strip() for v in row_values):
                                sheet_text.append(" | ".join(row_values))
                        if len(sheet_text) > 1:
                            extracted_text.append("\n".join(sheet_text))
                    file_content_text = "\n\n".join(extracted_text) if extracted_text else "[Excel without extractable data]"
                except Exception as excel_error:
                    file_content_text = f"[Error extracting Excel data: {str(excel_error)}]"

            elif file_extension == 'csv':
                try:
                    file_content_text = file_content.decode('utf-8')
                except UnicodeDecodeError:
                    file_content_text = file_content.decode('latin-1')

            elif file_extension in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp']:
                image_base64 = base64.b64encode(file_content).decode('utf-8')
                if PYTESSERACT_AVAILABLE:
                    try:
                        image = Image.open(io.BytesIO(file_content))
                        ocr_text = pytesseract.image_to_string(image)
                        file_content_text = ocr_text.strip() if ocr_text.strip() else "[Image without detected text]"
                    except Exception as ocr_error:
                        file_content_text = f"[OCR error: {str(ocr_error)}]"
                else:
                    file_content_text = "[Image file]"

            else:
                try:
                    file_content_text = file_content.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        file_content_text = file_content.decode('latin-1')
                    except:
                        file_content_text = f"[Binary file: {file_name}]"

        except Exception as e:
            file_content_text = f"[Error reading file: {str(e)}]"

    # Need either prompt or file
    if not prompt_text and not uploaded_file:
        return JsonResponse({'error': 'Please enter a message or attach a file.'}, status=400)

    # Build the full content for the message
    if file_content_text and prompt_text:
        full_content = f"{prompt_text}\n\n--- Attached File: {file_name} ---\n{file_content_text[:10000]}"
        display_content = prompt_text
    elif file_content_text:
        full_content = f"Please analyze this file:\n\n--- File: {file_name} ---\n{file_content_text[:10000]}"
        display_content = f"[Attached file: {file_name}]"
        prompt_text = display_content
    else:
        full_content = prompt_text
        display_content = prompt_text

    # Create new conversation if needed
    if not conversation:
        title = prompt_text[:50] if prompt_text else file_name[:50] if file_name else "New conversation"
        conversation = Conversation.objects.create(
            user=user,
            title=title + ('...' if len(title) > 50 else '')
        )

    # Create user message
    user_message = Message.objects.create(
        conversation=conversation,
        role='user',
        content=display_content,
        file_name=file_name,
        file_size=file_size
    )

    # Create prompt for audit/tracking
    prompt_obj = Prompt.objects.create(
        user=user,
        conversation=conversation,
        content=prompt_text or f"[File: {file_name}]",
        file_name=file_name,
        file_size=file_size,
        file_md5=file_hashes.get('md5') if file_hashes else None,
        file_sha1=file_hashes.get('sha1') if file_hashes else None,
        file_sha256=file_hashes.get('sha256') if file_hashes else None
    )

    # Check if user already confirmed consent
    consent_confirmed = request.POST.get('consent_confirmed') == 'true'

    # Validate prompt against rules
    requires_consent = False
    consent_rules_data = []
    try:
        prompt_matches, prompt_matched_rules = validate_prompt(user, prompt_text, prompt_obj, None, 'prompts')

        # Validate file content against rules
        file_matches = []
        file_matched_rules = []
        if file_content_text:
            file_info = {'name': file_name, 'size': file_size, 'hashes': file_hashes} if file_hashes else None
            file_matches, file_matched_rules = validate_prompt(user, file_content_text, prompt_obj, None, 'files', file_info)

        # Combine matches
        all_matches = (prompt_matches or []) + (file_matches or [])
        all_matched_rules = prompt_matched_rules + file_matched_rules

        # Separate rules by action type
        blocking_rules = [r for r in all_matched_rules if r.get('action', 'block') == 'block']
        consent_rules = [r for r in all_matched_rules if r.get('action') == 'consent']
        blocked_rules_data = []

        # Block rules always take priority
        if blocking_rules:
            for rule_data in blocking_rules:
                try:
                    from front.models import Rule
                    rule = Rule.objects.get(id=rule_data['id'])

                    Alert.objects.create(
                        user=user,
                        company=company,
                        prompt=prompt_obj,
                        rule=rule,
                        severity=rule_data.get('severity', 'medium'),
                        description=f"Prompt blocked by rule '{rule.name}'. The message was not sent to the LLM."
                    )

                    log_action(user, 'blocking_alert_created', {
                        'rule_name': rule.name,
                        'severity': rule_data.get('severity', 'medium'),
                        'prompt_id': prompt_obj.id,
                        'conversation_id': conversation.id
                    }, request)

                except Rule.DoesNotExist:
                    import logging
                    logging.warning(f"Rule {rule_data.get('name', 'unknown')} not found for alert creation")

            # Mark as blocked
            prompt_obj.filtered = True
            prompt_obj.save()
            user_message.blocked = True
            user_message.save()
            blocked = True
            blocked_rules_data = [{'name': r['name'], 'severity': r['severity'], 'description': r['description']} for r in blocking_rules]

        # Consent rules: ask for confirmation (only if no blocking rules)
        elif consent_rules and not consent_confirmed:
            requires_consent = True
            consent_rules_data = [{'name': r['name'], 'severity': r['severity'], 'description': r['description']} for r in consent_rules]

        # User confirmed consent: create consent alerts and proceed
        elif consent_rules and consent_confirmed:
            for rule_data in consent_rules:
                try:
                    from front.models import Rule
                    rule = Rule.objects.get(id=rule_data['id'])

                    Alert.objects.create(
                        user=user,
                        company=company,
                        prompt=prompt_obj,
                        rule=rule,
                        severity=rule_data.get('severity', 'medium'),
                        description=f"User consented to send message despite rule '{rule.name}' match."
                    )

                    log_action(user, 'consent_alert_created', {
                        'rule_name': rule.name,
                        'severity': rule_data.get('severity', 'medium'),
                        'prompt_id': prompt_obj.id,
                        'conversation_id': conversation.id
                    }, request)

                except Rule.DoesNotExist:
                    import logging
                    logging.warning(f"Rule {rule_data.get('name', 'unknown')} not found for alert creation")

    except Exception as validation_error:
        # Log validation error but don't block the request
        import logging
        logging.error(f"Error validating prompt: {validation_error}")
        all_matches = []
        blocked_rules_data = []

    # Log the action
    log_action(user, 'prompt_sent', {
        'prompt_id': prompt_obj.id,
        'conversation_id': conversation.id,
        'content_length': len(prompt_text) if prompt_text else 0,
        'has_file': bool(uploaded_file),
        'file_name': file_name,
        'file_size': file_size,
        'rules_matched': len(all_matches) if all_matches else 0,
        'blocked': blocked
    }, request)

    response_data = {
        'success': True,
        'conversation_id': conversation.id,
        'blocked': blocked,
        'blocked_rules': blocked_rules_data,
        'requires_consent': requires_consent,
        'consent_rules': consent_rules_data,
        'user_message': {
            'id': user_message.id,
            'content': display_content,
            'file_name': file_name,
            'file_size': file_size,
            'blocked': blocked
        }
    }

    # If consent is required, return early without sending to LLM
    if requires_consent:
        return JsonResponse(response_data)

    # Only send to LLM if not blocked
    if not blocked:
        try:
            # Build messages history for context using ChatMessage objects
            chat_messages = []
            for msg in conversation.messages.exclude(id=user_message.id):
                if msg.role in ['user', 'assistant']:
                    chat_messages.append(ChatMessage(
                        role=msg.role,
                        content=msg.content
                    ))

            # Add current message - handle images with vision API
            if is_image and image_base64:
                user_content = []
                if prompt_text and prompt_text != f"[Attached file: {file_name}]":
                    user_content.append({
                        "type": "text",
                        "text": prompt_text
                    })
                else:
                    user_content.append({
                        "type": "text",
                        "text": "Please analyze this image."
                    })
                user_content.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{image_mime_type};base64,{image_base64}"
                    }
                })
                chat_messages.append(ChatMessage(
                    role="user",
                    content=user_content
                ))
            else:
                chat_messages.append(ChatMessage(
                    role="user",
                    content=full_content
                ))

            # Use connector to send message
            ai_response = connector.chat(
                messages=chat_messages,
                model=model_name,
                max_tokens=1000,
                temperature=0.6,
            )

            # Save assistant response as message
            assistant_message = Message.objects.create(
                conversation=conversation,
                role='assistant',
                content=ai_response.content,
                tokens_used=ai_response.tokens_used
            )

            # Update conversation timestamp
            conversation.save()

            response_data['assistant_message'] = {
                'id': assistant_message.id,
                'content': ai_response.content
            }

        except Exception as e:
            error_str = str(e).lower()
            # Detect authentication/API key errors
            is_auth_error = any(keyword in error_str for keyword in [
                'authentication', 'auth', 'api key', 'api_key', 'apikey',
                'unauthorized', '401', 'invalid_api_key', 'invalid api key',
                'incorrect api key', 'invalid x-api-key',
            ])

            if is_auth_error:
                response_data['error_type'] = 'invalid_api_key'
                response_data['error'] = str(e)
                response_data['engine_name'] = current_engine.name if current_engine else 'Default'
                # Save error as system message
                error_message = Message.objects.create(
                    conversation=conversation,
                    role='system',
                    content=f"Error: Invalid API key for the configured AI engine."
                )
                response_data['error_message'] = {
                    'id': error_message.id,
                    'content': f"Error: Invalid API key for the configured AI engine."
                }
            else:
                # Save error as system message
                error_message = Message.objects.create(
                    conversation=conversation,
                    role='system',
                    content=f"Error: {str(e)}"
                )
                response_data['error_message'] = {
                    'id': error_message.id,
                    'content': f"Error: {str(e)}"
                }

    return JsonResponse(response_data)


def validate_prompt(user, prompt_text, prompt_instance=None, file=None, content_type='prompts', file_info=None):
    """
    Validates a prompt or file content against YARA rules in the DB.

    :param user: User sending the prompt
    :param prompt_text: Prompt text or file content
    :param prompt_instance: (optional) related Prompt object, to link in alerts
    :param file: (optional) uploaded file (not used currently)
    :param content_type: 'prompts' or 'files' - type of content being validated
    :param file_info: (optional) dict with file info including hashes {'name', 'size', 'hashes': {'md5', 'sha1', 'sha256'}}
    :return: tuple (list of matches, list of rules that matched)
    """
    if not prompt_text:
        return [], []

    # 1. Get active rules from company and user that apply to this content type
    # Rules apply if: applies_to == content_type OR applies_to == 'both'
    valid_applies_to = [content_type, 'both']

    import logging
    logging.info(f"validate_prompt called: content_type={content_type}, text_length={len(prompt_text) if prompt_text else 0}")

    company_rules = []
    if hasattr(user, 'profile') and user.profile.company:
        company_rules = Rule.objects.filter(
            rules_group__company=user.profile.company,
            active=True,
            applies_to__in=valid_applies_to
        )
        logging.info(f"Found {company_rules.count()} company rules for content_type={content_type}")

    user_rules = Rule.objects.filter(
        rules_group__user=user,
        active=True,
        applies_to__in=valid_applies_to
    )
    logging.info(f"Found {user_rules.count()} user rules for content_type={content_type}")

    all_rules = list(company_rules) + list(user_rules)

    if not all_rules:
        logging.info(f"No rules found for content_type={content_type}")
        return [], []  # no rules, no validation

    logging.info(f"Total rules to validate: {len(all_rules)}")

    # 2. Prepare dictionary to compile all rules
    rules_dict = {}
    for r in all_rules:
        rules_dict[f"rule_{r.id}"] = r.yara_rule

    try:
        compiled_rules = yara.compile(sources=rules_dict)
    except Exception as e:
        raise ValueError(f"Error compiling YARA rules: {e}")

    # 3. Execute match against the prompt
    matches = compiled_rules.match(data=prompt_text)
    logging.info(f"YARA matching completed: {len(matches)} matches found")

    # 4. Process matches and gather rule information
    # Create a dictionary of rules by ID for quick access
    rules_by_id = {r.id: r for r in all_rules}

    # List to store information about rules that matched
    matched_rules_info = []

    for match in matches:
        # Get the rule ID and its severity
        rule_id = int(match.namespace.split('_')[1])
        rule = rules_by_id.get(rule_id)
        severity = rule.severity if rule else "medium"

        # Save rule information to display in the modal
        matched_rules_info.append({
            'id': rule.id if rule else None,
            'name': rule.name if rule else 'Unknown Rule',
            'description': rule.description if rule else '',
            'severity': severity,
            'action': rule.action if rule else 'block',
        })

    # NOTE: No creamos alertas aquí - se crean en chat_send_message según el tipo de regla
    # (bloqueo vs ofuscación)

    return matches, matched_rules_info


# ==================== RULES GROUPS VIEWS ====================

@security_required
def rulesgroups_list(request):
    """Lists all rules groups for the user's company only."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Filter groups: only user's company groups (excluding global/community rules)
    groups = RulesGroup.objects.filter(
        company=user_company
    ).select_related('company', 'user').annotate(
        rules_count=Count('rules'),
        active_rules_count=Count('rules', filter=models.Q(rules__active=True))
    ).order_by('-created_at')

    # Filters
    search = request.GET.get('search', '')

    if search:
        groups = groups.filter(name__icontains=search)

    # Statistics - only for user's company
    total_groups = RulesGroup.objects.filter(company=user_company).count()
    total_rules = Rule.objects.filter(rules_group__company=user_company).count()
    active_rules = Rule.objects.filter(rules_group__company=user_company, active=True).count()

    context = {
        'groups': groups,
        'total_groups': total_groups,
        'total_rules': total_rules,
        'active_rules': active_rules,
        'search': search,
        'user_company': user_company,
    }
    return render(request, 'rulesgroups/list.html', context)


@security_required
def rulesgroup_detail(request, pk):
    """Shows the details of a rules group."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow access to company groups (not global/community)
    group = get_object_or_404(
        RulesGroup.objects.filter(
            company=user_company
        ).select_related('company', 'user').annotate(
            rules_count=Count('rules'),
            active_rules_count=Count('rules', filter=models.Q(rules__active=True))
        ),
        pk=pk
    )

    # Group rules
    rules = Rule.objects.filter(rules_group=group).order_by('-created_at')

    context = {
        'group': group,
        'rules': rules,
        'can_edit': True,  # User can always edit their company's groups
    }
    return render(request, 'rulesgroups/detail.html', context)


@security_required
def rulesgroup_create(request):
    """Creates a new rules group."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Users only from user's company
    users = User.objects.filter(profile__company=user_company)

    if request.method == 'POST':
        name = request.POST.get('name')
        description = request.POST.get('description', '')

        # Always create group for user's company
        group = RulesGroup.objects.create(
            name=name,
            description=description,
            company=user_company,
            user=None,  # No user-specific groups, only company-wide
        )
        log_action(request.user, 'rulesgroup_create', {
            'group_id': group.id,
            'group_name': group.name,
            'company_id': user_company.id if user_company else None,
        }, request)
        messages.success(request, f'Group "{group.name}" created successfully.')
        return redirect('front:rulesgroup_detail', pk=group.pk)

    return render(request, 'rulesgroups/form.html', {
        'user_company': user_company,
    })


@security_required
def rulesgroup_edit(request, pk):
    """Edits an existing rules group."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow editing company groups (not global)
    group = get_object_or_404(RulesGroup, pk=pk, company=user_company)

    if request.method == 'POST':
        group.name = request.POST.get('name')
        group.description = request.POST.get('description', '')
        # Keep company as is, don't allow changing it
        group.user = None  # No user-specific groups
        group.save()

        log_action(request.user, 'rulesgroup_update', {
            'group_id': group.id,
            'group_name': group.name
        }, request)
        messages.success(request, f'Group "{group.name}" updated successfully.')
        return redirect('front:rulesgroup_detail', pk=group.pk)

    return render(request, 'rulesgroups/form.html', {
        'group': group,
        'user_company': user_company,
    })


@security_required
def rulesgroup_delete(request, pk):
    """Deletes a rules group."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow deleting company groups (not global)
    group = get_object_or_404(RulesGroup, pk=pk, company=user_company)

    if request.method == 'POST':
        name = group.name
        group_id = group.id
        rules_count = Rule.objects.filter(rules_group=group).count()
        group.delete()
        log_action(request.user, 'rulesgroup_delete', {
            'group_id': group_id,
            'group_name': name,
            'rules_deleted': rules_count
        }, request)
        messages.success(request, f'Group "{name}" deleted successfully.')
        return redirect('front:rulesgroups_list')

    # Count rules that will be deleted
    rules_count = Rule.objects.filter(rules_group=group).count()

    return render(request, 'rulesgroups/delete.html', {
        'group': group,
        'rules_count': rules_count,
    })


@security_required
def rulesgroup_activate_all(request, pk):
    """Activates all rules in a rules group."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow activating rules in company groups (not global)
    group = get_object_or_404(RulesGroup, pk=pk, company=user_company)

    if request.method == 'POST':
        count = Rule.objects.filter(rules_group=group, active=False).update(active=True)
        log_action(request.user, 'rulesgroup_activate_all', {
            'group_id': group.id,
            'group_name': group.name,
            'rules_activated': count
        }, request)
        messages.success(request, f'{count} rules activated in group "{group.name}".')

    next_url = request.GET.get('next', 'front:rulesgroup_detail')
    if next_url.startswith('/'):
        return redirect(next_url)
    return redirect('front:rulesgroup_detail', pk=pk)


@security_required
def rulesgroup_deactivate_all(request, pk):
    """Deactivates all rules in a rules group."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow deactivating rules in company groups (not global)
    group = get_object_or_404(RulesGroup, pk=pk, company=user_company)

    if request.method == 'POST':
        count = Rule.objects.filter(rules_group=group, active=True).update(active=False)
        log_action(request.user, 'rulesgroup_deactivate_all', {
            'group_id': group.id,
            'group_name': group.name,
            'rules_deactivated': count
        }, request)
        messages.warning(request, f'{count} rules deactivated in group "{group.name}".')

    next_url = request.GET.get('next', 'front:rulesgroup_detail')
    if next_url.startswith('/'):
        return redirect(next_url)
    return redirect('front:rulesgroup_detail', pk=pk)


# ==================== RULES VIEWS ====================

@security_required
def rules_list(request):
    """Lists all YARA rules for the user's company only."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Filter rules: only company rules (excluding global/community)
    rules = Rule.objects.filter(
        rules_group__company=user_company
    ).select_related('rules_group', 'rules_group__company').order_by('-created_at')

    # Filters
    status_filter = request.GET.get('status')
    group_filter = request.GET.get('group')
    search = request.GET.get('search', '')

    if status_filter == 'active':
        rules = rules.filter(active=True)
    elif status_filter == 'inactive':
        rules = rules.filter(active=False)

    if group_filter:
        rules = rules.filter(rules_group_id=group_filter)

    if search:
        rules = rules.filter(name__icontains=search)

    # Statistics - only for user's company
    base_rules = Rule.objects.filter(rules_group__company=user_company)
    total_rules = base_rules.count()
    active_rules = base_rules.filter(active=True).count()
    inactive_rules = base_rules.filter(active=False).count()

    # Groups for the filter - only company groups
    rules_groups = RulesGroup.objects.filter(company=user_company)

    context = {
        'rules': rules,
        'rules_groups': rules_groups,
        'total_rules': total_rules,
        'active_rules': active_rules,
        'inactive_rules': inactive_rules,
        'current_status': status_filter,
        'current_group': group_filter,
        'search': search,
        'user_company': user_company,
    }
    return render(request, 'rules/list.html', context)


@security_required
def rule_detail(request, pk):
    """Shows the details of a YARA rule."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow access to company rules (not global/community)
    rule = get_object_or_404(
        Rule.objects.filter(
            rules_group__company=user_company
        ).select_related('rules_group', 'rules_group__company'),
        pk=pk
    )

    # Alerts generated by this rule
    alerts = Alert.objects.filter(rule=rule, company=user_company).select_related('user', 'prompt').order_by('-created_at')[:10]

    context = {
        'rule': rule,
        'alerts': alerts,
        'alerts_count': Alert.objects.filter(rule=rule, company=user_company).count(),
        'can_edit': True,  # User can always edit their company's rules
    }
    return render(request, 'rules/detail.html', context)


@security_required
def rule_create(request):
    """Creates a new YARA rule."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only show company groups (not global) for creating new rules
    rules_groups = RulesGroup.objects.filter(company=user_company)

    if request.method == 'POST':
        name = request.POST.get('name')
        description = request.POST.get('description', '')
        yara_rule = request.POST.get('yara_rule')
        rules_group_id = request.POST.get('rules_group')
        severity = request.POST.get('severity', 'medium')
        applies_to = request.POST.get('applies_to', 'both')
        action = request.POST.get('action', 'block')
        active = request.POST.get('active') == 'on'

        # Validate required fields
        if not name or not yara_rule or not rules_group_id:
            messages.error(request, 'Please fill in all required fields.')
            return render(request, 'rules/form.html', {
                'rules_groups': rules_groups,
                'form_data': request.POST,
            })

        # Validate that the group belongs to user's company
        group = RulesGroup.objects.filter(pk=rules_group_id, company=user_company).first()
        if not group:
            messages.error(request, 'Invalid rules group selected.')
            return render(request, 'rules/form.html', {
                'rules_groups': rules_groups,
                'form_data': request.POST,
            })

        # Validate YARA rule
        try:
            yara.compile(source=yara_rule)
        except yara.SyntaxError as e:
            messages.error(request, f'YARA syntax error: {e}')
            return render(request, 'rules/form.html', {
                'rules_groups': rules_groups,
                'form_data': request.POST,
            })
        except Exception as e:
            messages.error(request, f'YARA validation error: {e}')
            return render(request, 'rules/form.html', {
                'rules_groups': rules_groups,
                'form_data': request.POST,
            })

        try:
            rule = Rule.objects.create(
                name=name,
                description=description,
                yara_rule=yara_rule,
                rules_group_id=rules_group_id,
                severity=severity,
                applies_to=applies_to,
                action=action,
                active=active,
            )
            log_action(request.user, 'rule_create', {
                'rule_id': rule.id,
                'rule_name': rule.name,
                'severity': severity,
                'applies_to': applies_to,
                'action': action,
                'active': active
            }, request)
            messages.success(request, f'Rule "{rule.name}" created successfully.')
            return redirect('front:rules_list')
        except Exception as e:
            messages.error(request, f'Error creating rule: {e}')
            return render(request, 'rules/form.html', {
                'rules_groups': rules_groups,
                'form_data': request.POST,
            })

    return render(request, 'rules/form.html', {
        'rules_groups': rules_groups,
    })


@security_required
def rule_edit(request, pk):
    """Edits an existing YARA rule."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow editing company rules (not global)
    rule = get_object_or_404(Rule, pk=pk, rules_group__company=user_company)

    # Only show company groups for editing
    rules_groups = RulesGroup.objects.filter(company=user_company)

    if request.method == 'POST':
        rule.name = request.POST.get('name')
        rule.description = request.POST.get('description', '')
        rule.yara_rule = request.POST.get('yara_rule')
        rules_group_id = request.POST.get('rules_group')
        rule.severity = request.POST.get('severity', 'medium')
        rule.applies_to = request.POST.get('applies_to', 'both')
        rule.action = request.POST.get('action', 'block')
        rule.active = request.POST.get('active') == 'on'

        # Validate required fields
        if not rule.name or not rule.yara_rule or not rules_group_id:
            messages.error(request, 'Please fill in all required fields.')
            return render(request, 'rules/form.html', {
                'rule': rule,
                'rules_groups': rules_groups,
            })

        # Validate that the group belongs to user's company
        group = RulesGroup.objects.filter(pk=rules_group_id, company=user_company).first()
        if not group:
            messages.error(request, 'Invalid rules group selected.')
            return render(request, 'rules/form.html', {
                'rule': rule,
                'rules_groups': rules_groups,
            })

        rule.rules_group_id = rules_group_id

        # Validate YARA rule
        try:
            yara.compile(source=rule.yara_rule)
        except yara.SyntaxError as e:
            messages.error(request, f'YARA syntax error: {e}')
            return render(request, 'rules/form.html', {
                'rule': rule,
                'rules_groups': rules_groups,
            })
        except Exception as e:
            messages.error(request, f'YARA validation error: {e}')
            return render(request, 'rules/form.html', {
                'rule': rule,
                'rules_groups': rules_groups,
            })

        try:
            rule.save()
            log_action(request.user, 'rule_update', {
                'rule_id': rule.id,
                'rule_name': rule.name,
                'severity': rule.severity,
                'applies_to': rule.applies_to,
                'action': rule.action,
                'active': rule.active
            }, request)
            messages.success(request, f'Rule "{rule.name}" updated successfully.')
            return redirect('front:rule_detail', pk=rule.pk)
        except Exception as e:
            messages.error(request, f'Error updating rule: {e}')
            return render(request, 'rules/form.html', {
                'rule': rule,
                'rules_groups': rules_groups,
            })

    return render(request, 'rules/form.html', {
        'rule': rule,
        'rules_groups': rules_groups,
    })


@security_required
def rule_delete(request, pk):
    """Deletes a YARA rule."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow deleting company rules (not global)
    rule = get_object_or_404(Rule, pk=pk, rules_group__company=user_company)

    if request.method == 'POST':
        name = rule.name
        rule_id = rule.id
        rule.delete()
        log_action(request.user, 'rule_delete', {
            'rule_id': rule_id,
            'rule_name': name
        }, request)
        messages.success(request, f'Rule "{name}" deleted successfully.')
        return redirect('front:rules_list')

    return render(request, 'rules/delete.html', {'rule': rule})


@security_required
def rule_toggle(request, pk):
    """Activates/deactivates a YARA rule."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow toggling company rules (not global)
    rule = get_object_or_404(Rule, pk=pk, rules_group__company=user_company)
    rule.active = not rule.active
    rule.save()

    log_action(request.user, 'rule_toggle', {
        'rule_id': rule.id,
        'rule_name': rule.name,
        'new_status': 'active' if rule.active else 'inactive'
    }, request)

    status = 'activated' if rule.active else 'deactivated'
    messages.success(request, f'Rule "{rule.name}" {status}.')

    # Redirect to previous page or list
    next_url = request.GET.get('next', 'front:rules_list')
    if next_url.startswith('/'):
        return redirect(next_url)
    return redirect(next_url)


# ==================== RULES COMMUNITY VIEWS ====================

@security_required
def rules_community_list(request):
    """Lists all community (global) rules available for configuration."""
    user_company = get_user_company(request.user)

    # Get global rules (rules in groups with no company and no user)
    community_rules = Rule.objects.filter(
        rules_group__company__isnull=True,
        rules_group__user__isnull=True
    ).select_related('rules_group').order_by('rules_group__name', 'name')

    # Get IDs of rules already configured for this company
    # (rules that were duplicated from community rules)
    configured_rule_names = Rule.objects.filter(
        rules_group__company=user_company
    ).values_list('name', flat=True)

    # Mark which rules are already configured
    for rule in community_rules:
        rule.is_configured = rule.name in configured_rule_names

    # Filters
    search = request.GET.get('search', '')
    severity_filter = request.GET.get('severity')
    applies_to_filter = request.GET.get('applies_to')
    status_filter = request.GET.get('status')  # configured/not_configured

    if search:
        community_rules = [r for r in community_rules if search.lower() in r.name.lower() or (r.description and search.lower() in r.description.lower())]

    if severity_filter:
        community_rules = [r for r in community_rules if r.severity == severity_filter]

    if applies_to_filter:
        community_rules = [r for r in community_rules if r.applies_to == applies_to_filter]

    if status_filter == 'configured':
        community_rules = [r for r in community_rules if r.is_configured]
    elif status_filter == 'not_configured':
        community_rules = [r for r in community_rules if not r.is_configured]

    # Group rules by their RulesGroup for better display
    groups_dict = {}
    for rule in community_rules:
        group_name = rule.rules_group.name
        if group_name not in groups_dict:
            groups_dict[group_name] = {
                'group': rule.rules_group,
                'rules': []
            }
        groups_dict[group_name]['rules'].append(rule)

    # Statistics
    total_community_rules = Rule.objects.filter(
        rules_group__company__isnull=True,
        rules_group__user__isnull=True
    ).count()
    configured_count = len([r for r in community_rules if r.is_configured])

    context = {
        'groups_dict': groups_dict,
        'community_rules': community_rules,
        'total_community_rules': total_community_rules,
        'configured_count': configured_count,
        'search': search,
        'severity_filter': severity_filter,
        'applies_to_filter': applies_to_filter,
        'status_filter': status_filter,
    }
    return render(request, 'rules_community/list.html', context)


@security_required
def rule_community_configure(request, pk):
    """Configures (duplicates) a community rule for the user's company."""
    user_company = get_user_company(request.user)

    # Get the community rule
    community_rule = get_object_or_404(
        Rule,
        pk=pk,
        rules_group__company__isnull=True,
        rules_group__user__isnull=True
    )

    # Check if already configured
    existing = Rule.objects.filter(
        rules_group__company=user_company,
        name=community_rule.name
    ).first()

    if existing:
        return JsonResponse({
            'success': False,
            'error': 'already_configured',
            'message': f'Rule "{community_rule.name}" is already configured for your company.',
            'rule_id': existing.id
        })

    # Find or create a RulesGroup for community rules in this company
    company_group, created = RulesGroup.objects.get_or_create(
        company=user_company,
        name=f"Community Rules",
        defaults={
            'description': 'Rules imported from the community repository'
        }
    )

    # Duplicate the rule for the company (inactive by default)
    new_rule = Rule.objects.create(
        name=community_rule.name,
        description=community_rule.description,
        yara_rule=community_rule.yara_rule,
        rules_group=company_group,
        severity=community_rule.severity,
        applies_to=community_rule.applies_to,
        active=False  # Inactive by default, user must activate
    )

    log_action(request.user, 'rule_community_configure', {
        'community_rule_id': community_rule.id,
        'community_rule_name': community_rule.name,
        'new_rule_id': new_rule.id,
        'company': user_company.name
    }, request)

    return JsonResponse({
        'success': True,
        'message': f'Rule "{community_rule.name}" has been configured for your company.',
        'rule_id': new_rule.id,
        'rule_name': new_rule.name,
        'group_name': company_group.name
    })


@security_required
def rule_community_activate(request, pk):
    """Activates a recently configured community rule."""
    user_company = get_user_company(request.user)

    rule = get_object_or_404(Rule, pk=pk, rules_group__company=user_company)
    rule.active = True
    rule.save()

    log_action(request.user, 'rule_activate', {
        'rule_id': rule.id,
        'rule_name': rule.name
    }, request)

    return JsonResponse({
        'success': True,
        'message': f'Rule "{rule.name}" has been activated.',
        'rule_id': rule.id
    })


@admin_required
def rules_community_sync(request):
    """
    Syncs community rules from the GitHub repository.
    Only admins can sync rules. Rules are stored in a global RulesGroup (company=null, user=null).
    """
    GITHUB_REPO = "Regulai/regulai-rules"
    GITHUB_API_URL = f"https://api.github.com/repos/{GITHUB_REPO}/contents"
    GITHUB_RAW_URL = f"https://raw.githubusercontent.com/{GITHUB_REPO}/main"

    try:
        # Get list of files from GitHub
        response = requests.get(GITHUB_API_URL, timeout=10)
        response.raise_for_status()
        files = response.json()

        # Get or create the global "Community" RulesGroup
        community_group, created = RulesGroup.objects.get_or_create(
            company__isnull=True,
            user__isnull=True,
            name="Community",
            defaults={
                'description': 'Community-maintained YARA rules from the regulai-rules repository'
            }
        )

        rules_added = 0
        rules_updated = 0
        rules_skipped = 0
        errors = []

        for file_info in files:
            if file_info['type'] != 'file':
                continue

            file_name = file_info['name']
            download_url = file_info.get('download_url')

            if not download_url:
                continue

            try:
                # Download the file content
                file_response = requests.get(download_url, timeout=10)
                file_response.raise_for_status()
                file_content = file_response.text

                # Parse YARA rules from the file
                parsed_rules = parse_yara_file(file_content, file_name)

                for rule_data in parsed_rules:
                    # Check if rule already exists
                    existing_rule = Rule.objects.filter(
                        rules_group=community_group,
                        name=rule_data['name']
                    ).first()

                    if existing_rule:
                        # Update if content changed
                        if existing_rule.yara_rule != rule_data['yara_rule']:
                            existing_rule.yara_rule = rule_data['yara_rule']
                            existing_rule.description = rule_data.get('description', '')
                            existing_rule.severity = rule_data.get('severity', 'medium')
                            existing_rule.save()
                            rules_updated += 1
                        else:
                            rules_skipped += 1
                    else:
                        # Create new rule
                        Rule.objects.create(
                            name=rule_data['name'],
                            description=rule_data.get('description', ''),
                            yara_rule=rule_data['yara_rule'],
                            rules_group=community_group,
                            severity=rule_data.get('severity', 'medium'),
                            applies_to='both',
                            active=True  # Community rules are active by default
                        )
                        rules_added += 1

            except Exception as e:
                errors.append(f"Error processing {file_name}: {str(e)}")

        log_action(request.user, 'rules_community_sync', {
            'rules_added': rules_added,
            'rules_updated': rules_updated,
            'rules_skipped': rules_skipped,
            'errors': errors
        }, request)

        return JsonResponse({
            'success': True,
            'message': f'Sync completed: {rules_added} added, {rules_updated} updated, {rules_skipped} unchanged.',
            'rules_added': rules_added,
            'rules_updated': rules_updated,
            'rules_skipped': rules_skipped,
            'errors': errors if errors else None
        })

    except requests.RequestException as e:
        return JsonResponse({
            'success': False,
            'message': f'Error connecting to GitHub: {str(e)}'
        }, status=500)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'message': f'Error during sync: {str(e)}'
        }, status=500)


def parse_yara_file(content, filename):
    """
    Parses a YARA file and extracts individual rules with their metadata.
    Returns a list of dictionaries with rule information.
    """
    rules = []

    # Regex to match YARA rules
    # Matches: rule RuleName : tags { ... }
    rule_pattern = re.compile(
        r'rule\s+(\w+)\s*(?::\s*[\w\s]+)?\s*\{(.*?)\n\}',
        re.DOTALL
    )

    # Find all rules in the file
    for match in rule_pattern.finditer(content):
        rule_name = match.group(1)
        rule_body = match.group(0)  # Full rule including 'rule name { ... }'

        # Extract metadata
        meta = {}
        meta_match = re.search(r'meta:\s*(.*?)(?=strings:|condition:|$)', rule_body, re.DOTALL)
        if meta_match:
            meta_content = meta_match.group(1)
            # Parse key = "value" pairs
            for meta_item in re.finditer(r'(\w+)\s*=\s*"([^"]*)"', meta_content):
                meta[meta_item.group(1).lower()] = meta_item.group(2)

        rules.append({
            'name': rule_name,
            'description': meta.get('description', meta.get('purpose', f'Rule from {filename}')),
            'severity': meta.get('severity', 'medium'),
            'yara_rule': rule_body.strip()
        })

    return rules


# ==================== ALERTS VIEWS ====================

@security_required
def alerts_list(request):
    """Lists all alerts with filters."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Filter alerts by company
    alerts = Alert.objects.filter(
        company=user_company
    ).select_related('user', 'user__profile', 'user__profile__department', 'company', 'prompt', 'rule').order_by('-created_at')

    # Filters
    severity_filter = request.GET.get('severity')
    status_filter = request.GET.get('status')
    rule_filter = request.GET.get('rule')
    department_filter = request.GET.get('department')
    search = request.GET.get('search', '')

    if severity_filter:
        alerts = alerts.filter(severity=severity_filter)

    if status_filter == 'resolved':
        alerts = alerts.filter(resolved=True)
    elif status_filter == 'unresolved':
        alerts = alerts.filter(resolved=False)

    if rule_filter:
        alerts = alerts.filter(rule_id=rule_filter)

    if department_filter:
        alerts = alerts.filter(user__profile__department_id=department_filter)

    if search:
        alerts = alerts.filter(description__icontains=search)

    # Statistics - filtered by company
    total_alerts = Alert.objects.filter(company=user_company).count()
    unresolved_alerts = Alert.objects.filter(company=user_company, resolved=False).count()
    critical_alerts = Alert.objects.filter(company=user_company, severity='critical', resolved=False).count()
    high_alerts = Alert.objects.filter(company=user_company, severity='high', resolved=False).count()
    medium_alerts = Alert.objects.filter(company=user_company, severity='medium', resolved=False).count()
    low_alerts = Alert.objects.filter(company=user_company, severity='low', resolved=False).count()

    # Alerts by severity for chart - filtered by company
    alerts_by_severity = Alert.objects.filter(company=user_company, resolved=False).values('severity').annotate(count=Count('id'))

    # Rules for the filter - company rules + global rules that have alerts
    rules = Rule.objects.filter(
        Q(rules_group__company=user_company) | Q(rules_group__company__isnull=True),
        alert__company=user_company
    ).distinct()

    # Departments for the filter - only from user's company
    departments = Department.objects.filter(company=user_company)

    context = {
        'alerts': alerts,
        'rules': rules,
        'departments': departments,
        'total_alerts': total_alerts,
        'unresolved_alerts': unresolved_alerts,
        'critical_alerts': critical_alerts,
        'high_alerts': high_alerts,
        'medium_alerts': medium_alerts,
        'low_alerts': low_alerts,
        'alerts_by_severity': alerts_by_severity,
        'current_severity': severity_filter,
        'current_status': status_filter,
        'current_rule': rule_filter,
        'current_department': department_filter,
        'search': search,
    }
    return render(request, 'alerts/list.html', context)


@security_required
def alert_detail(request, pk):
    """Shows the details of an alert."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow access to alerts from user's company
    alert = get_object_or_404(
        Alert.objects.select_related('user', 'company', 'prompt', 'rule', 'rule__rules_group'),
        pk=pk,
        company=user_company
    )

    # Get the response associated with the prompt if it exists
    response = None
    if alert.prompt:
        try:
            response = Response.objects.get(prompt=alert.prompt)
        except Response.DoesNotExist:
            pass

    context = {
        'alert': alert,
        'response': response,
    }
    return render(request, 'alerts/detail.html', context)


@security_required
def alert_resolve(request, pk):
    """Marks an alert as resolved."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow access to alerts from user's company
    alert = get_object_or_404(Alert, pk=pk, company=user_company)
    alert.resolved = True
    alert.save()

    log_action(request.user, 'alert_resolve', {
        'alert_id': alert.id,
        'severity': alert.severity,
        'rule_name': alert.rule.name if alert.rule else None
    }, request)

    messages.success(request, f'Alert #{alert.id} marked as resolved.')

    next_url = request.GET.get('next', 'front:alerts_list')
    if next_url.startswith('/'):
        return redirect(next_url)
    return redirect(next_url)


@security_required
def alert_unresolve(request, pk):
    """Marks an alert as unresolved."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow access to alerts from user's company
    alert = get_object_or_404(Alert, pk=pk, company=user_company)
    alert.resolved = False
    alert.save()

    log_action(request.user, 'alert_unresolve', {
        'alert_id': alert.id,
        'severity': alert.severity,
        'rule_name': alert.rule.name if alert.rule else None
    }, request)

    messages.warning(request, f'Alert #{alert.id} marked as unresolved.')

    next_url = request.GET.get('next', 'front:alerts_list')
    if next_url.startswith('/'):
        return redirect(next_url)
    return redirect(next_url)


@security_required
def alert_delete(request, pk):
    """Deletes an alert."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow access to alerts from user's company
    alert = get_object_or_404(Alert, pk=pk, company=user_company)

    if request.method == 'POST':
        alert_id = alert.id
        severity = alert.severity
        rule_name = alert.rule.name if alert.rule else None
        alert.delete()
        log_action(request.user, 'alert_delete', {
            'alert_id': alert_id,
            'severity': severity,
            'rule_name': rule_name
        }, request)
        messages.success(request, f'Alert #{alert_id} deleted successfully.')
        return redirect('front:alerts_list')

    return render(request, 'alerts/delete.html', {'alert': alert})


@security_required
def alerts_resolve_all(request):
    """Marks all unresolved alerts as resolved."""
    if request.method == 'POST':
        # Get user's company
        user_company = get_user_company(request.user)

        # Only resolve alerts from user's company
        count = Alert.objects.filter(company=user_company, resolved=False).update(resolved=True)
        log_action(request.user, 'alerts_resolve_all', {
            'alerts_resolved': count
        }, request)
        messages.success(request, f'{count} alerts marked as resolved.')

    return redirect('front:alerts_list')


# ==================== AUDIT LOGS VIEWS ====================

@admin_required
def auditlogs_list(request):
    """Lists all audit logs with filters."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Filter logs by company (through user's profile) or system logs (user=None)
    company_or_system = Q(user__profile__company=user_company) | Q(user__isnull=True)
    logs = AuditLog.objects.filter(
        company_or_system
    ).select_related('user').order_by('-timestamp')

    # Filters
    action_filter = request.GET.get('action')
    user_filter = request.GET.get('user')
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    search = request.GET.get('search', '')

    if action_filter:
        logs = logs.filter(action=action_filter)

    if user_filter:
        if user_filter == 'system':
            logs = logs.filter(user__isnull=True)
        else:
            logs = logs.filter(user_id=user_filter)

    if date_from:
        logs = logs.filter(timestamp__date__gte=date_from)

    if date_to:
        logs = logs.filter(timestamp__date__lte=date_to)

    if search:
        logs = logs.filter(details__icontains=search)

    # Statistics - filtered by company or system
    total_logs = AuditLog.objects.filter(company_or_system).count()
    today_logs = AuditLog.objects.filter(
        company_or_system,
        timestamp__date=timezone.now().date()
    ).count()

    # Unique actions for the filter - from company's logs and system logs
    actions = AuditLog.objects.filter(
        company_or_system
    ).values_list('action', flat=True).distinct()

    # Users for the filter - only from user's company
    users = User.objects.filter(profile__company=user_company, auditlog__isnull=False).distinct()

    # Logs by action for the chart - filtered by company and system
    logs_by_action = AuditLog.objects.filter(
        company_or_system
    ).values('action').annotate(count=Count('id')).order_by('-count')[:5]

    context = {
        'logs': logs[:100],  # Limitar a 100 registros
        'actions': actions,
        'users': users,
        'total_logs': total_logs,
        'today_logs': today_logs,
        'logs_by_action': logs_by_action,
        'current_action': action_filter,
        'current_user': user_filter,
        'date_from': date_from,
        'date_to': date_to,
        'search': search,
    }
    return render(request, 'auditlogs/list.html', context)


@admin_required
def auditlog_detail(request, pk):
    """Shows the details of an audit log entry."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow access to logs from user's company or system logs
    company_or_system = Q(user__profile__company=user_company) | Q(user__isnull=True)
    log = get_object_or_404(
        AuditLog.objects.filter(company_or_system).select_related('user'),
        pk=pk,
    )

    context = {
        'log': log,
    }
    return render(request, 'auditlogs/detail.html', context)


@admin_required
def auditlogs_clear(request):
    """Deletes all old audit log entries (more than 30 days)."""
    if request.method == 'POST':
        # Get user's company
        user_company = get_user_company(request.user)

        cutoff_date = timezone.now() - timezone.timedelta(days=30)
        # Delete logs from user's company or system logs
        company_or_system = Q(user__profile__company=user_company) | Q(user__isnull=True)
        count = AuditLog.objects.filter(
            company_or_system,
            timestamp__lt=cutoff_date
        ).delete()[0]
        messages.success(request, f'{count} old records deleted.')

    return redirect('front:auditlogs_list')


@admin_required
def auditlog_delete(request, pk):
    """Deletes an audit log entry."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow access to logs from user's company or system logs
    company_or_system = Q(user__profile__company=user_company) | Q(user__isnull=True)
    log = get_object_or_404(AuditLog.objects.filter(company_or_system), pk=pk)

    if request.method == 'POST':
        log_id = log.id
        log.delete()
        messages.success(request, f'Record #{log_id} deleted successfully.')
        return redirect('front:auditlogs_list')

    return render(request, 'auditlogs/delete.html', {'log': log})


# ==================== PROMPTS VIEWS ====================

@security_required
def prompts_list(request):
    """Lists all executed prompts with filters."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Filter prompts by company (through user's profile)
    prompts = Prompt.objects.filter(
        user__profile__company=user_company
    ).select_related('user', 'user__profile', 'user__profile__department').order_by('-created_at')

    # Filters
    user_filter = request.GET.get('user')
    model_filter = request.GET.get('model')
    filtered_filter = request.GET.get('filtered')
    department_filter = request.GET.get('department')
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    search = request.GET.get('search', '')

    if user_filter:
        prompts = prompts.filter(user_id=user_filter)

    if model_filter:
        prompts = prompts.filter(model_used=model_filter)

    if filtered_filter == 'yes':
        prompts = prompts.filter(filtered=True)
    elif filtered_filter == 'no':
        prompts = prompts.filter(filtered=False)

    if department_filter:
        prompts = prompts.filter(user__profile__department_id=department_filter)

    if date_from:
        prompts = prompts.filter(created_at__date__gte=date_from)

    if date_to:
        prompts = prompts.filter(created_at__date__lte=date_to)

    if search:
        prompts = prompts.filter(content__icontains=search)

    # Statistics - filtered by company
    total_prompts = Prompt.objects.filter(user__profile__company=user_company).count()
    today_prompts = Prompt.objects.filter(
        user__profile__company=user_company,
        created_at__date=timezone.now().date()
    ).count()
    filtered_prompts = Prompt.objects.filter(
        user__profile__company=user_company,
        filtered=True
    ).count()

    # Users for the filter - only from user's company
    users = User.objects.filter(profile__company=user_company, prompt__isnull=False).distinct()

    # Departments for the filter - only from user's company
    departments = Department.objects.filter(company=user_company)

    # Unique models for the filter - from company's prompts
    models = Prompt.objects.filter(
        user__profile__company=user_company
    ).values_list('model_used', flat=True).distinct()

    # Prompts per day (last 7 days) - filtered by company
    prompts_by_day = Prompt.objects.filter(
        user__profile__company=user_company,
        created_at__gte=timezone.now() - timezone.timedelta(days=7)
    ).annotate(
        date=TruncDate('created_at')
    ).values('date').annotate(count=Count('id')).order_by('date')

    context = {
        'prompts': prompts[:100],  # Limitar a 100 registros
        'users': users,
        'departments': departments,
        'models': models,
        'total_prompts': total_prompts,
        'today_prompts': today_prompts,
        'filtered_prompts': filtered_prompts,
        'prompts_by_day': prompts_by_day,
        'current_user': user_filter,
        'current_model': model_filter,
        'current_filtered': filtered_filter,
        'current_department': department_filter,
        'date_from': date_from,
        'date_to': date_to,
        'search': search,
    }
    return render(request, 'prompts/list.html', context)


@security_required
def prompt_detail(request, pk):
    """Shows the details of a prompt."""
    # Get user's company
    user_company = get_user_company(request.user)

    # Only allow access to prompts from user's company
    prompt = get_object_or_404(
        Prompt.objects.select_related('user'),
        pk=pk,
        user__profile__company=user_company
    )

    # Get the associated response if it exists
    response = None
    try:
        response = Response.objects.get(prompt=prompt)
    except Response.DoesNotExist:
        pass

    # Get alerts generated by this prompt
    alerts = Alert.objects.filter(prompt=prompt).select_related('rule')

    context = {
        'prompt': prompt,
        'response': response,
        'alerts': alerts,
    }
    return render(request, 'prompts/detail.html', context)

